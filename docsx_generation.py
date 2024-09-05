from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

class DOCXReport:
    def __init__(self, company_name):
        self.company_name = company_name
        self.document = Document()
        self.document.add_heading(f"FX Risk Rating Analysis for {self.company_name}", level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add company information
        self.document.add_paragraph(f"Company: {self.company_name}").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self.document.add_paragraph("Report Date: August 2024").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self.document.add_paragraph()

    def add_chapter(self, title, body):
        self.document.add_heading(title, level=2)
        
        # Process body text similar to how it is done in the PDF function
        for line in body.splitlines():
            if line.startswith("### "):
                self.add_section_title(line.replace("### ", "").strip())
            elif line.startswith("#### "):
                self.add_sub_section_title(line.replace("#### ", "").strip())
            else:
                self.add_body_text(line)

    def add_section_title(self, title):
        self.document.add_heading(title, level=3)

    def add_sub_section_title(self, title):
        self.document.add_heading(title, level=4)

    def add_body_text(self, text):
        # Replace Markdown-style bold markers and add text
        text = text.replace("**", "")
        self.document.add_paragraph(text)

    def add_bullet_point(self, text):
        paragraph = self.document.add_paragraph(style='List Bullet')
        run = paragraph.add_run(u"\u2022 " + text)
        run.font.size = Pt(12)

    def add_numbered_item(self, number, text):
        paragraph = self.document.add_paragraph(style='List Number')
        run = paragraph.add_run(f"{number}. {text}")
        run.font.size = Pt(12)

    def save(self, filename):
        output_directory = "LLM_result"
        os.makedirs(output_directory, exist_ok=True)

        docx_path = os.path.join(output_directory, filename)
        self.document.save(docx_path)
        print(f"Document saved at {docx_path}")

def save_output_to_docx(title, analysis_data, company_name, output_path):
    # Initialize the DOCX report object with the company name
    docx_report = DOCXReport(company_name)
    docx_report.add_chapter(title, analysis_data)
    
    # Save the DOCX to the specified output path
    docx_report.save(output_path)
    
    print(f"DOCX saved at {output_path}")
